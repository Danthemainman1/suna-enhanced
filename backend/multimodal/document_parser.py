"""
Document parsing with multiple format support.

This module provides parsing capabilities for PDF, DOCX, XLSX,
and other document formats.
"""

import uuid
from typing import Optional
from pathlib import Path
from .models import ParsedDocument, Table


class DocumentParser:
    """
    Multi-format document parser.
    
    Parses PDF, DOCX, XLSX, CSV, and Markdown documents,
    extracting text, tables, and metadata.
    """
    
    def __init__(self):
        """Initialize the document parser."""
        pass
    
    async def parse(
        self,
        file: bytes,
        filename: str
    ) -> ParsedDocument:
        """
        Parse a document and extract content.
        
        Args:
            file: Document file bytes
            filename: Original filename
            
        Returns:
            ParsedDocument with extracted content
        """
        # Determine file type from extension
        ext = Path(filename).suffix.lower()
        
        if ext == '.pdf':
            return await self._parse_pdf(file, filename)
        elif ext in ['.docx', '.doc']:
            return await self._parse_docx(file, filename)
        elif ext in ['.xlsx', '.xls']:
            return await self._parse_xlsx(file, filename)
        elif ext == '.csv':
            return await self._parse_csv(file, filename)
        elif ext in ['.md', '.markdown']:
            return await self._parse_markdown(file, filename)
        else:
            # Try to parse as plain text
            return await self._parse_text(file, filename)
    
    async def _parse_pdf(
        self,
        file: bytes,
        filename: str
    ) -> ParsedDocument:
        """Parse PDF document."""
        try:
            from PyPDF2 import PdfReader
            import io
            
            reader = PdfReader(io.BytesIO(file))
            
            # Extract metadata
            metadata = reader.metadata if reader.metadata else {}
            title = metadata.get('/Title', None)
            author = metadata.get('/Author', None)
            
            # Extract text from all pages
            text_content = ""
            for page in reader.pages:
                text_content += page.extract_text() + "\n"
            
            return ParsedDocument(
                filename=filename,
                document_type='pdf',
                title=title,
                author=author,
                text_content=text_content.strip(),
                num_pages=len(reader.pages),
                metadata={
                    'producer': metadata.get('/Producer'),
                    'subject': metadata.get('/Subject'),
                }
            )
        except Exception as e:
            # Fallback to basic parsing
            return ParsedDocument(
                filename=filename,
                document_type='pdf',
                text_content=f"Error parsing PDF: {str(e)}",
                metadata={'error': str(e)}
            )
    
    async def _parse_docx(
        self,
        file: bytes,
        filename: str
    ) -> ParsedDocument:
        """Parse DOCX document."""
        try:
            from docx import Document
            import io
            
            doc = Document(io.BytesIO(file))
            
            # Extract text from paragraphs
            text_content = "\n".join([para.text for para in doc.paragraphs])
            
            # Extract tables
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = self._extract_docx_table(table, i)
                tables.append(table_data)
            
            # Try to get core properties
            try:
                title = doc.core_properties.title
                author = doc.core_properties.author
                created = doc.core_properties.created
            except:
                title = None
                author = None
                created = None
            
            return ParsedDocument(
                filename=filename,
                document_type='docx',
                title=title,
                author=author,
                created_date=created,
                text_content=text_content.strip(),
                num_pages=1,  # DOCX doesn't have page concept
                tables=tables
            )
        except Exception as e:
            return ParsedDocument(
                filename=filename,
                document_type='docx',
                text_content=f"Error parsing DOCX: {str(e)}",
                metadata={'error': str(e)}
            )
    
    def _extract_docx_table(self, table, table_index: int) -> Table:
        """Extract table from DOCX."""
        rows_data = []
        headers = []
        
        for i, row in enumerate(table.rows):
            row_data = [cell.text for cell in row.cells]
            if i == 0:
                headers = row_data
            else:
                rows_data.append(row_data)
        
        return Table(
            table_id=f"table_{table_index}",
            headers=headers,
            rows=rows_data
        )
    
    async def _parse_xlsx(
        self,
        file: bytes,
        filename: str
    ) -> ParsedDocument:
        """Parse XLSX document."""
        try:
            from openpyxl import load_workbook
            import io
            
            wb = load_workbook(io.BytesIO(file))
            
            # Extract text and tables from all sheets
            text_content = ""
            tables = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_content += f"\n\n=== {sheet_name} ===\n\n"
                
                # Get all rows with data
                rows_data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        rows_data.append([str(cell) if cell is not None else "" for cell in row])
                
                if rows_data:
                    # First row as headers
                    headers = rows_data[0] if rows_data else []
                    data_rows = rows_data[1:] if len(rows_data) > 1 else []
                    
                    table = Table(
                        table_id=f"sheet_{sheet_name}",
                        headers=headers,
                        rows=data_rows
                    )
                    tables.append(table)
                    
                    # Add to text content
                    for row in rows_data:
                        text_content += " | ".join(row) + "\n"
            
            return ParsedDocument(
                filename=filename,
                document_type='xlsx',
                text_content=text_content.strip(),
                num_pages=len(wb.sheetnames),
                tables=tables
            )
        except Exception as e:
            return ParsedDocument(
                filename=filename,
                document_type='xlsx',
                text_content=f"Error parsing XLSX: {str(e)}",
                metadata={'error': str(e)}
            )
    
    async def _parse_csv(
        self,
        file: bytes,
        filename: str
    ) -> ParsedDocument:
        """Parse CSV file."""
        try:
            import csv
            import io
            
            # Try to decode as UTF-8
            try:
                content = file.decode('utf-8')
            except UnicodeDecodeError:
                # Try other encodings
                import chardet
                encoding = chardet.detect(file)['encoding']
                content = file.decode(encoding)
            
            reader = csv.reader(io.StringIO(content))
            rows_data = list(reader)
            
            # First row as headers
            headers = rows_data[0] if rows_data else []
            data_rows = rows_data[1:] if len(rows_data) > 1 else []
            
            table = Table(
                table_id="csv_table",
                headers=headers,
                rows=data_rows
            )
            
            # Create text representation
            text_content = "\n".join([" | ".join(row) for row in rows_data])
            
            return ParsedDocument(
                filename=filename,
                document_type='csv',
                text_content=text_content,
                num_pages=1,
                tables=[table]
            )
        except Exception as e:
            return ParsedDocument(
                filename=filename,
                document_type='csv',
                text_content=f"Error parsing CSV: {str(e)}",
                metadata={'error': str(e)}
            )
    
    async def _parse_markdown(
        self,
        file: bytes,
        filename: str
    ) -> ParsedDocument:
        """Parse Markdown file."""
        try:
            content = file.decode('utf-8')
            
            # Extract title (first # heading)
            title = None
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    title = line[2:].strip()
                    break
            
            return ParsedDocument(
                filename=filename,
                document_type='markdown',
                title=title,
                text_content=content,
                num_pages=1
            )
        except Exception as e:
            return ParsedDocument(
                filename=filename,
                document_type='markdown',
                text_content=f"Error parsing Markdown: {str(e)}",
                metadata={'error': str(e)}
            )
    
    async def _parse_text(
        self,
        file: bytes,
        filename: str
    ) -> ParsedDocument:
        """Parse plain text file."""
        try:
            # Try UTF-8 first
            try:
                content = file.decode('utf-8')
            except UnicodeDecodeError:
                # Try to detect encoding
                import chardet
                encoding = chardet.detect(file)['encoding']
                content = file.decode(encoding or 'latin-1')
            
            return ParsedDocument(
                filename=filename,
                document_type='text',
                text_content=content,
                num_pages=1
            )
        except Exception as e:
            return ParsedDocument(
                filename=filename,
                document_type='text',
                text_content=f"Error parsing text: {str(e)}",
                metadata={'error': str(e)}
            )
    
    async def extract_tables(
        self,
        file: bytes,
        filename: str
    ) -> list[Table]:
        """
        Extract only tables from a document.
        
        Args:
            file: Document file bytes
            filename: Original filename
            
        Returns:
            List of extracted tables
        """
        doc = await self.parse(file, filename)
        return doc.tables
    
    async def to_markdown(
        self,
        file: bytes,
        filename: str
    ) -> str:
        """
        Convert document to Markdown format.
        
        Args:
            file: Document file bytes
            filename: Original filename
            
        Returns:
            Markdown formatted text
        """
        doc = await self.parse(file, filename)
        
        markdown = ""
        
        # Add title
        if doc.title:
            markdown += f"# {doc.title}\n\n"
        
        # Add metadata
        if doc.author:
            markdown += f"**Author:** {doc.author}\n\n"
        
        # Add content
        markdown += doc.text_content + "\n\n"
        
        # Add tables
        for table in doc.tables:
            markdown += "\n"
            # Headers
            if table.headers:
                markdown += "| " + " | ".join(table.headers) + " |\n"
                markdown += "| " + " | ".join(["---"] * len(table.headers)) + " |\n"
            # Rows
            for row in table.rows:
                markdown += "| " + " | ".join(row) + " |\n"
            markdown += "\n"
        
        return markdown
